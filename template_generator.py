# template_generator.py
# Generate FIR and RTI templates (in DOCX) for the Nyay Buddy app

from docx import Document
from io import BytesIO
import datetime

# -------------------------------------------------------
# FIR TEMPLATE GENERATOR
# -------------------------------------------------------
def generate_fir(data: dict) -> bytes:
    """Generate a First Information Report (FIR) document and return as bytes."""
    doc = Document()
    doc.add_heading("FIRST INFORMATION REPORT (FIR)", level=1)
    doc.add_paragraph(f"Date: {datetime.date.today().strftime('%d-%m-%Y')}")
    doc.add_paragraph("To,")
    doc.add_paragraph("The Station House Officer,\nNearest Police Station\nCity / District")

    doc.add_paragraph("\nSubject: Lodging of FIR regarding incident")

    doc.add_paragraph(f"Sir/Madam,\n\nI, {data.get('complainant_name','[Your Name]')}, "
                      f"residing at {data.get('complainant_address','[Your Address]')}, "
                      f"wish to lodge an FIR regarding the incident that occurred on "
                      f"{data.get('incident_date','[Date]')} at {data.get('incident_place','[Place]')}.")

    doc.add_paragraph(f"\nDetails of Incident:\n{data.get('incident_details','[Incident details]')}")

    if data.get("accused_name"):
        doc.add_paragraph(f"\nAccused Person(s): {data['accused_name']}")

    if data.get("suspected_sections"):
        doc.add_paragraph(f"\nRelevant Section(s): {data['suspected_sections']}")

    doc.add_paragraph("\nI request you to kindly register the FIR and take appropriate legal action as per law.")
    doc.add_paragraph("\nThanking You,\nYours faithfully,\n")
    doc.add_paragraph(f"Name: {data.get('complainant_name','')}")
    doc.add_paragraph(f"Contact: {data.get('contact','')}")
    doc.add_paragraph("\n(Signature)\n")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# -------------------------------------------------------
# RTI TEMPLATE GENERATOR
# -------------------------------------------------------
def generate_rti_application(data: dict) -> bytes:
    """Generate RTI Application in DOCX format."""
    doc = Document()
    doc.add_heading("RIGHT TO INFORMATION (RTI) APPLICATION", level=1)
    doc.add_paragraph(f"Date: {datetime.date.today().strftime('%d-%m-%Y')}")
    doc.add_paragraph("To,\nThe Public Information Officer (PIO)\n[Department Name]\n[Address]\n\nSubject: Request for Information under RTI Act, 2005")

    doc.add_paragraph(f"Sir/Madam,\n\nI, {data.get('applicant_name','[Your Name]')}, "
                      f"residing at {data.get('address','[Your Address]')}, "
                      f"wish to obtain the following information under the RTI Act, 2005:\n")

    doc.add_paragraph(f"1. {data.get('information_requested','[Describe information required]')}")

    doc.add_paragraph("\nI have paid the prescribed application fee of ₹10 as per RTI Rules.\n\nKindly provide the requested information within 30 days as per Section 7(1) of the RTI Act, 2005.")

    doc.add_paragraph("\nThanking You,\nYours faithfully,\n")
    doc.add_paragraph(f"Name: {data.get('applicant_name','')}")
    doc.add_paragraph(f"Address: {data.get('address','')}")
    doc.add_paragraph(f"Contact: {data.get('contact','')}")
    doc.add_paragraph("\n(Signature)\n")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()
